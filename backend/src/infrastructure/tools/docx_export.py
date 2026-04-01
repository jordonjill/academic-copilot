"""
DOCX export tool.
"""
from __future__ import annotations

import os
from pathlib import Path

from langchain_core.tools import tool


def _resolve_output_path(output_path: str) -> Path:
    base_dir = Path(os.getenv("EXPORT_BASE_DIR", "data/exports")).expanduser()
    base_dir = base_dir.resolve()

    candidate = Path(output_path).expanduser()
    if not candidate.is_absolute():
        candidate = base_dir / candidate

    resolved = candidate.resolve()
    if resolved != base_dir and base_dir not in resolved.parents:
        raise ValueError("Output path must be within EXPORT_BASE_DIR.")

    try:
        resolved.parent.mkdir(parents=True, exist_ok=True)
    except OSError as exc:
        raise ValueError(f"Cannot create export directory '{resolved.parent}': {exc!r}") from exc
    return resolved


@tool
def export_docx(title: str, content: str, output_path: str) -> dict[str, str]:
    """Export content to a DOCX file and return its path."""
    from docx import Document

    path = _resolve_output_path(output_path)

    document = Document()
    if title:
        document.add_heading(title, level=1)
    for line in content.splitlines() or [""]:
        document.add_paragraph(line)

    try:
        document.save(path)
    except OSError as exc:
        raise ValueError(f"Cannot write DOCX file '{path}': {exc!r}") from exc
    return {"path": str(path)}
